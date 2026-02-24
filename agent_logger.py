"""
Agent Logger - DOCX logging for all agent calls.

Records input, model name, and output for every agent invocation.
Supports embedding generated code and figures (from base64) into the document.
"""

import base64
import io
import json
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class AgentLogger:
    """Logs all agent calls to a DOCX file."""

    def __init__(self, output_dir: str = "outputs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.output_path = self.output_dir / f"agent_log_{timestamp}.docx"

        self.doc = Document()
        self._setup_document()
        self._call_counter = 0

    def _setup_document(self):
        """Add title and metadata to the document."""
        title = self.doc.add_heading("Agent Call Log", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        self.doc.add_paragraph()  # spacer

    def log_agent_call(
        self,
        agent_name: str,
        model_name: str,
        input_data: Any,
        output_data: Any,
    ):
        """Log a standard agent call (no code/figure).

        Args:
            agent_name: Name of the agent (e.g., "Evaluator", "Supervisor").
            model_name: Model identifier (e.g., "claude-sonnet-4-5").
            input_data: The changeable input passed to the agent (not prompts).
            output_data: The output produced by the agent.
        """
        self._call_counter += 1

        # Heading
        self.doc.add_heading(
            f"{self._call_counter}. {agent_name}  [{model_name}]", level=2
        )

        # Input
        self.doc.add_heading("Input", level=3)
        input_text = self._format_data(input_data)
        p = self.doc.add_paragraph()
        run = p.add_run(input_text)
        run.font.size = Pt(9)

        # Output
        self.doc.add_heading("Output", level=3)
        output_text = self._format_data(output_data)
        p = self.doc.add_paragraph()
        run = p.add_run(output_text)
        run.font.size = Pt(9)

        self.doc.add_paragraph()  # spacer
        self._autosave()

    def log_code_and_figure(
        self,
        agent_name: str,
        model_name: str,
        input_data: Any,
        code: str,
        figure_base64: Optional[str] = None,
        extra_output: Any = None,
    ):
        """Log a code-generation agent call, embedding the code and figure.

        Args:
            agent_name: Name of the agent.
            model_name: Model identifier.
            input_data: The changeable input.
            code: The generated Python code.
            figure_base64: Base64-encoded PNG image (optional).
            extra_output: Any additional output data (e.g., summary_data).
        """
        self._call_counter += 1

        # Heading
        self.doc.add_heading(
            f"{self._call_counter}. {agent_name}  [{model_name}]", level=2
        )

        # Input
        self.doc.add_heading("Input", level=3)
        input_text = self._format_data(input_data)
        p = self.doc.add_paragraph()
        run = p.add_run(input_text)
        run.font.size = Pt(9)

        # Code
        self.doc.add_heading("Generated Code", level=3)
        p = self.doc.add_paragraph()
        run = p.add_run(code)
        run.font.size = Pt(8)
        run.font.name = "Courier New"

        # Figure
        if figure_base64:
            self.doc.add_heading("Generated Figure", level=3)
            try:
                image_bytes = base64.b64decode(figure_base64)
                image_stream = io.BytesIO(image_bytes)
                self.doc.add_picture(image_stream, width=Inches(6))
            except Exception as e:
                self.doc.add_paragraph(f"[Error embedding figure: {e}]")

        # Extra output
        if extra_output is not None:
            self.doc.add_heading("Additional Output", level=3)
            p = self.doc.add_paragraph()
            run = p.add_run(self._format_data(extra_output))
            run.font.size = Pt(9)

        self.doc.add_paragraph()  # spacer
        self._autosave()

    def log_section_header(self, title: str):
        """Add a section header to visually separate phases."""
        self.doc.add_page_break()
        heading = self.doc.add_heading(title, level=1)
        self.doc.add_paragraph()

    def log_user_message(self, message_index: int, message: str):
        """Log a user message as a subsection."""
        self.doc.add_heading(
            f"User Message #{message_index}: {message[:80]}{'...' if len(message) > 80 else ''}",
            level=2,
        )
        p = self.doc.add_paragraph()
        run = p.add_run(message)
        run.font.size = Pt(10)
        run.bold = True
        self.doc.add_paragraph()

    def save(self):
        """Save the DOCX file."""
        self.doc.save(str(self.output_path))
        print(f"DOCX saved to: {self.output_path}")

    def _autosave(self):
        """Auto-save after each log entry to avoid data loss."""
        try:
            self.doc.save(str(self.output_path))
        except Exception:
            pass  # non-critical

    def _format_data(self, data: Any) -> str:
        """Convert data to a readable string."""
        if isinstance(data, str):
            return data
        if isinstance(data, dict) or isinstance(data, list):
            try:
                return json.dumps(data, indent=2, ensure_ascii=False, default=str)
            except Exception:
                return str(data)
        return str(data)
